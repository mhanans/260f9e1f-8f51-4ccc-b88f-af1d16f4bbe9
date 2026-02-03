from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
import io
import docx
import structlog

logger = structlog.get_logger()

class OCREngine:
    def extract_text_from_image(self, image_bytes: bytes, filename: str = "image") -> str:
        try:
            image = Image.open(io.BytesIO(image_bytes))
            text = pytesseract.image_to_string(image)
            logger.info(f"OCR | Image {filename} | Extracted {len(text)} chars", preview=text[:50].replace('\n', ' '))
            return text
        except Exception as e:
            logger.error(f"OCR Error (Image) - {filename}: {e}")
            return ""

    def extract_text_from_pdf(self, pdf_bytes: bytes, filename: str = "doc.pdf") -> str:
        try:
            # Need to specify popler path if on windows usually, but assuming env is ok
            images = convert_from_bytes(pdf_bytes)
            full_text = ""
            logger.info(f"OCR | PDF {filename} | Found {len(images)} pages")
            
            for i, img in enumerate(images):
                page_text = pytesseract.image_to_string(img)
                full_text += page_text + "\n"
                logger.info(f"OCR | PDF {filename} | Page {i+1} | Extracted {len(page_text)} chars")
            
            return full_text
        except Exception as e:
            logger.error(f"OCR Error (PDF) - {filename}: {e}")
            return ""

    def extract_text_from_docx(self, docx_bytes: bytes, filename: str = "doc.docx") -> str:
        try:
            doc = docx.Document(io.BytesIO(docx_bytes))
            full_text = "\n".join([p.text for p in doc.paragraphs])
            
            logger.info(f"OCR | DOCX {filename} | Text Paragraphs: {len(doc.paragraphs)}")
            
            # Extract images from DOCX
            image_count = 0
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    image_count += 1
                    try:
                        img_bytes = rel.target_part.blob
                        img_text = self.extract_text_from_image(img_bytes, f"{filename}_img_{image_count}")
                        full_text += f"\n[Image {image_count} Content]: " + img_text
                    except Exception as inner:
                        logger.error(f"OCR Error (DOCX Image {image_count}) - {filename}: {inner}")
            
            logger.info(f"OCR | DOCX {filename} | Found {image_count} Embedded Images | Total Text Length: {len(full_text)}")
            return full_text
        except Exception as e:
            logger.error(f"OCR Error (DOCX) - {filename}: {e}")
            return ""

ocr_engine = OCREngine()
